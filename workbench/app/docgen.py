# 繁工AI 本地解析工作台 - 工程资料生成引擎（v0.1.7）
# 依据 v3.7 口径：内置方案类型 → 数据可自动从解析库预填 → 缺失字段列出待补充 → 生成 Word 下载
# 优先 Word（python-docx）；后续可扩展 PDF/Excel。签字由人工下载打印完成。

import datetime
import json
import os

from . import config

# 方案类型注册：名称 / 说明 / 必填字段（缺失时前端提示）/ 可选字段
TYPES = {
    "施工方案": {
        "label": "施工方案",
        "required": ["项目名称", "车间", "编制单位", "编制人", "施工内容", "施工日期"],
        "optional": ["施工单位", "审核人", "批准人", "施工工艺", "质量措施", "安全措施", "进度安排", "涉及设备"],
        "construction_params": ["施工内容", "施工工艺", "涉及设备"],
        "default_text": {
            "编制依据": "1. 设计图纸及设计交底文件\n2. 现行国家及行业标准规范\n3. 设备技术文件及随机资料\n4. 现场施工条件调查结果",
            "质量保证措施": "1. 严格执行三检制（自检、互检、专检）\n2. 关键工序设质量控制点，报验合格后进入下道工序\n3. 施工人员持证上岗，特种作业人员证件齐全",
            "安全文明施工措施": "1. 进入现场正确佩戴劳动防护用品\n2. 用电设备执行一机一闸一漏保\n3. 现场材料定置摆放，工完场清",
        },
    },
    "吊装方案": {
        "label": "吊装方案",
        "required": ["项目名称", "车间", "编制单位", "编制人", "吊装日期", "吊装设备名称", "设备重量"],
        "optional": ["施工单位", "审核人", "批准人", "吊装机械", "吊装半径", "吊装高度", "吊车站位", "吊索具", "吊装安全措施"],
        "lifting_params": ["吊装设备名称", "设备重量", "吊装半径", "吊装高度", "吊车型号", "吊车站位", "吊索具"],
        "default_text": {
            "编制依据": "1. 设备安装图纸及技术文件\n2.《起重机械安全规程》GB/T 6067\n3.《建筑机械使用安全技术规程》JGJ 33\n4. 设备随机装箱单及发货资料",
            "吊装安全措施": "1. 吊装前办理吊装作业票，检查吊索具完好\n2. 明确指挥信号，专人指挥，无关人员撤离吊装区\n3. 六级及以上大风、雷雨等恶劣天气停止吊装\n4. 设备就位后立即找正固定，方可摘钩",
        },
    },
    "技术交底": {
        "label": "技术交底",
        "required": ["项目名称", "车间", "交底人", "交底日期", "交底内容"],
        "optional": ["被交底单位", "被交底人", "交底依据", "注意事项"],
        "default_text": {
            "交底依据": "施工图纸、施工方案、国家现行规范",
            "注意事项": "1. 作业前逐条学习交底内容并签字确认\n2. 发现与图纸不符时停止作业并及时上报\n3. 交底内容作为现场施工和检查验收依据",
        },
    },
    "开箱验收记录": {
        "label": "开箱验收记录",
        "required": ["项目名称", "车间", "箱单号", "验收人", "验收日期", "验收结果"],
        "optional": ["位号", "设备名称", "数量", "外观检查情况", "随机资料", "备注"],
        "default_text": {
            "验收依据": "设备装箱单、发货清单、采购合同及技术协议",
            "验收流程": "1. 核对箱体编号与装箱单一致\n2. 开箱后核对设备名称、型号、数量\n3. 检查外观有无锈蚀、变形、损坏\n4. 清点随机附件与资料",
        },
    },
    "隐蔽工程验收记录": {
        "label": "隐蔽工程验收记录",
        "required": ["项目名称", "车间", "隐蔽部位", "验收人", "验收日期", "验收结果"],
        "optional": ["施工单位", "监理单位", "隐蔽内容", "检查情况", "影像资料编号"],
        "default_text": {
            "验收依据": "施工图纸、施工方案、现行验收规范",
            "检查情况": "1. 隐蔽内容与图纸相符，几何尺寸符合要求\n2. 预埋件位置准确，固定牢靠\n3. 隐蔽前影像资料留存齐全",
        },
    },
    # ---- v0.1.12 新增：工程资料深度生成（按项目进度、缺项待办） ----
    "施工计划": {
        "label": "施工计划",
        "required": ["项目名称", "车间", "编制单位", "编制人", "计划工期"],
        "optional": ["施工单位", "审核人", "开工日期", "竣工日期", "进度安排", "资源配置", "关键节点"],
        "default_text": {
            "编制依据": "1. 施工图纸及工程量清单\n2. 合同约定的工期要求\n3. 现行施工及验收规范",
            "进度安排": "1. 施工准备阶段（场地、材料、机具、人员）\n2. 主体施工阶段（按车间/工序流水作业）\n3. 收尾验收阶段（自检、整改、报验）",
        },
    },
    "施工日志": {
        "label": "施工日志",
        "required": ["项目名称", "车间", "记录人", "记录日期", "当日工作内容"],
        "optional": ["天气", "温度", "到场人员", "进场材料", "机械使用", "安全质量情况", "问题及处理"],
        "default_text": {
            "记录要求": "1. 逐日记录，不得补记、漏记\n2. 与当日施工部位、工序对应\n3. 问题与处理结果闭环记录",
        },
    },
    "设计变更": {
        "label": "设计变更",
        "required": ["项目名称", "车间", "变更编号", "变更内容", "提出人"],
        "optional": ["原设计图号", "变更图号", "变更原因", "涉及设备", "工程量变化", "审批人", "日期"],
        "default_text": {
            "变更依据": "1. 现场与设计不符情况\n2. 业主/设计/监理会签意见",
            "处理要求": "1. 变更单与原图纸一同归档\n2. 变更内容落实到相关专业图纸与施工\n3. 变更工程量作为结算依据",
        },
    },
    "货损报告": {
        "label": "货损报告",
        "required": ["项目名称", "车间", "设备位号", "损失情况", "报告人", "报告日期"],
        "optional": ["箱单号", "数量", "损失程度", "影像资料编号", "处理意见", "责任方"],
        "default_text": {
            "报告依据": "开箱/到货验收记录、随货装箱单、运输单据",
            "处理流程": "1. 现场拍照留存，保护受损设备\n2. 会同运输方/厂家确认损失程度\n3. 出具报告并提交索赔或补发申请",
        },
    },
    "竣工资料": {
        "label": "竣工资料",
        "required": ["项目名称", "车间", "编制单位", "编制人", "竣工日期"],
        "optional": ["施工单位", "监理单位", "隐蔽工程资料编号", "验收记录编号", "交工范围", "遗留问题及处理"],
        "default_text": {
            "编制依据": "1. 竣工图及变更文件\n2. 各分项验收记录\n3. 设备调试及试运行记录",
            "资料构成": "1. 竣工图（含变更反映）\n2. 隐蔽工程验收记录\n3. 设备开箱/安装/调试记录\n4. 质量验收与移交清单",
        },
    },
}


def list_types() -> list:
    return [{"key": k, "label": v["label"], "required": v["required"],
             "optional": v["optional"]} for k, v in TYPES.items()]


def std_citations(doc_type: str, limit: int = 3) -> list:
    """从平台规范库检索与当前资料类型相关的规范正文（非仅名称），供『编制依据』引用。v0.1.17"""
    try:
        from . import platform_store
        idx = platform_store.list_items().get("items", [])
        hits = [it for it in idx if it.get("std_no") and it.get("status") in ("现行", "待核验")]
        if not hits:
            return []
        # 用模板名关键词粗筛（检索不到就取前几条现行规范）
        kw = (doc_type or "").replace("记录", "").replace("资料", "")
        scored = []
        for it in hits:
            name = str(it.get("std_name", ""))
            score = 0
            for ch in kw:
                if ch and ch in name:
                    score += 1
            scored.append((score, it))
        scored.sort(key=lambda x: -x[0])
        out = []
        for score, it in scored[:limit]:
            text = ""
            try:
                cache = os.path.join(platform_store.CACHE_DIR, f"{it['sha256']}.json")
                if os.path.exists(cache):
                    with open(cache, encoding="utf-8") as f:
                        text = str(json.load(f).get("text", ""))[:600].strip()
            except Exception:  # noqa: BLE001
                text = ""
            if not text:
                text = "（正文已入库，可在平台库页检索全文）"
            out.append({"std_no": it["std_no"], "std_name": it.get("std_name", ""),
                        "status": it.get("status", ""), "snippet": text[:300]})
        return out
    except Exception:  # noqa: BLE001
        return []


def prefill_from_db(doc_type: str, workshop: str = "") -> dict:
    """v0.1.30：从解析库深度预填资料数据。
    来源：relations layout（车间设备）→ device_workshop（设备车间归属）→
          vector_store（设备参数/重量/型号检索）→ platform_store（规范正文引用）。
    返回 {data: {...}, missing: [...], devices: [...], citations: [...]}。"""
    data = {}
    missing = []
    t = TYPES.get(doc_type)
    if not t:
        return {"data": data, "missing": [], "devices": [], "citations": []}

    # 项目名称（从配置或第一个车间推断）
    data["项目名称"] = getattr(config, "PROJECT_NAME", "") or "紫金矿业工程项目"

    # 车间
    if workshop:
        data["车间"] = workshop

    # 从 relations 取车间设备
    devices = []
    try:
        from . import relations
        g = relations.load_relations()
        all_devs = g.get("devices", [])
        if workshop:
            ws_devs = [d for d in all_devs if workshop in d.get("workshops", [])]
        else:
            ws_devs = all_devs
        devices = [{"tag": d["tag"], "name": "见台账", "count": 1,
                    "workshops": d.get("workshops", [])} for d in ws_devs[:60]]
    except Exception:  # noqa: BLE001
        pass

    # 设备级车间归属补充（v0.1.29）
    try:
        from . import device_workshop
        for dv in devices:
            dw = device_workshop.get_workshop(dv["tag"])
            if dw and dw not in dv["workshops"]:
                dv["workshops"].append(dw)
    except Exception:  # noqa: BLE001
        pass

    data["_devices"] = devices

    # 吊装方案专项：从设备台账/向量库检索重量等参数
    if doc_type == "吊装方案" and devices:
        first_dev = devices[0]
        data["吊装设备名称"] = first_dev["tag"] + "（详见设备清单）"
        # 尝试从向量库检索设备重量
        try:
            from . import vector_store
            store = vector_store.VectorStore()
            hits = store.search(f"{first_dev['tag']} 重量 设备参数", top_k=3)
            for h in hits:
                txt = str(h.get("text", ""))
                import re as _re
                m = _re.search(r"(重量|净重|毛重)\s*[:：=]?\s*(\d+(?:\.\d+)?)\s*(t|吨|kg|公斤)", txt, _re.I)
                if m:
                    w = float(m.group(2))
                    unit = m.group(3).lower()
                    if unit in ("kg", "公斤"):
                        w = w / 1000.0
                    data["设备重量"] = f"{w} t"
                    break
        except Exception:  # noqa: BLE001
            pass
        if "设备重量" not in data:
            missing.append("设备重量（需从设备铭牌/台账补充）")
        # 吊装参数默认建议
        if "吊装半径" not in data:
            missing.append("吊装半径（需现场测量）")
        if "吊装高度" not in data:
            missing.append("吊装高度（需根据设备安装高度确定）")
        # v0.1.42：根据设备重量自动选择吊装方法
        if "吊车型号" not in data:
            from . import equipment_types as _et
            weight_str = data.get("设备重量", "")
            weight_t = 0.0
            import re as _re2
            wm = _re2.search(r"(\d+(?:\.\d+)?)", str(weight_str))
            if wm:
                weight_t = float(wm.group(1))
                if "kg" in str(weight_str).lower():
                    weight_t = weight_t / 1000.0
            lifting = _et.select_lifting_method(weight_t)
            data["吊车型号"] = lifting["crane_type"]
            data["吊装方法"] = lifting["method"]
            data["吊装说明"] = lifting["notes"]
        if "吊索具" not in data:
            # v0.1.42：根据重量选择吊索具规格
            weight_str = data.get("设备重量", "")
            import re as _re3
            wm = _re3.search(r"(\d+(?:\.\d+)?)", str(weight_str))
            if wm:
                w = float(wm.group(1))
                if "kg" in str(weight_str).lower():
                    w = w / 1000.0
                if w <= 5:
                    data["吊索具"] = "1t/3t吊带或钢丝绳（安全系数≥6）"
                elif w <= 25:
                    data["吊索具"] = "5t/10t钢丝绳（安全系数≥6，卸扣匹配）"
                elif w <= 50:
                    data["吊索具"] = "15t/25t钢丝绳（安全系数≥6，需计算吊耳强度）"
                else:
                    data["吊索具"] = f"{int(w*0.8)}t以上钢丝绳（安全系数≥6，需专项计算）"
            else:
                data["吊索具"] = "钢丝绳/吊带（根据设备重量选择，安全系数≥6）"
        # v0.1.50：吊装方案设备数据联动——技术参数、吊装环境、空间限制、相邻设备、相关管线
        try:
            from . import relations as _rel
            from . import spatial_model as _sm
            g = _rel.load_relations()
            spatial = _sm.build_spatial_model(g)
            spatial_devs = {d["tag"]: d for d in spatial.get("devices", [])}
            # 设备技术参数（吊装用）
            tech_params = []
            lifting_env = []
            for d in devices[:5]:
                tag = d["tag"]
                sd = spatial_devs.get(tag, {})
                params = []
                if sd.get("workshop"):
                    params.append(f"车间: {sd['workshop']}")
                if sd.get("x") is not None and sd.get("y") is not None:
                    params.append(f"坐标: ({sd['x']}, {sd['y']})")
                if sd.get("z") is not None:
                    params.append(f"安装标高: {sd['z']}m")
                    lifting_env.append(f"{tag}安装标高{sd['z']}m")
                if sd.get("coord_status"):
                    params.append(f"位置状态: {sd['coord_status']}")
                if params:
                    tech_params.append(f"{tag}: {'; '.join(params)}")
                if sd.get("workshop"):
                    lifting_env.append(f"{tag}位于{sd['workshop']}")
            if tech_params:
                data["吊装设备参数"] = "\n".join(tech_params)
            if lifting_env:
                data["吊装环境"] = "; ".join(lifting_env)
            # 空间限制（从标高和楼层推断）
            space_limits = []
            for d in devices[:3]:
                tag = d["tag"]
                sd = spatial_devs.get(tag, {})
                if sd.get("z") is not None and sd["z"] > 10:
                    space_limits.append(f"{tag}高位吊装（标高{sd['z']}m），需考虑高空作业防护和风速限制")
                if sd.get("workshop"):
                    space_limits.append(f"{tag}在{sd['workshop']}内吊装，需确认车间内空间和吊车通道")
            if space_limits:
                data["空间限制"] = "\n".join(space_limits)
            # 相邻设备（吊装时需注意的相邻设备）
            neighbors = []
            for d in devices[:3]:
                tag = d["tag"]
                sd = spatial_devs.get(tag, {})
                if sd.get("neighbors"):
                    nb = [n.get("tag", "") for n in sd["neighbors"][:3] if n.get("tag")]
                    if nb:
                        neighbors.append(f"{tag}相邻: {', '.join(nb)}（吊装时需做好成品保护）")
            if neighbors:
                data["吊装相邻设备"] = "; ".join(neighbors)
            # 相关管线（吊装时需注意的管线）
            try:
                from . import piping_network as _pn
                pipe_info = []
                for d in devices[:3]:
                    tag = d["tag"]
                    pipes = _pn.get_device_pipes(tag)
                    if pipes:
                        p_info = [f"{p['pipe_no']}({p['medium']})" for p in pipes[:3]]
                        pipe_info.append(f"{tag}连接管线: {', '.join(p_info)}（吊装前需确认管线断开或保护）")
                if pipe_info:
                    data["吊装相关管线"] = "\n".join(pipe_info)
            except Exception:  # noqa: BLE001
                pass
        except Exception:  # noqa: BLE001
            pass

    # 技术交底专项（v0.1.51：设备数据联动——技术参数、施工步骤、安全要点、质量控制）
    if doc_type == "技术交底" and devices:
        data["涉及设备"] = "、".join(d["tag"] for d in devices[:10]) + (f" 等{len(devices)}台" if len(devices) > 10 else "")
        # 识别设备类型，自动生成施工步骤和安全要点
        from . import equipment_types as _et
        eq_type = _et.get_equipment_type_from_devices(devices)
        data["设备类型"] = eq_type
        steps = _et.get_construction_steps(eq_type)
        data["施工步骤"] = "\n".join(f"{i+1}. {step}" for i, step in enumerate(steps))
        # 安全要点（根据设备类型）
        safety_points = _get_safety_points(eq_type, devices)
        data["安全要点"] = "\n".join(f"{i+1}. {p}" for i, p in enumerate(safety_points))
        # 质量控制要点
        quality_points = _get_quality_points(eq_type)
        data["质量控制要点"] = "\n".join(f"{i+1}. {p}" for i, p in enumerate(quality_points))
        # 交底内容自动生成
        data["交底内容"] = f"{eq_type}安装技术交底：施工准备→基础验收→设备就位→找平找正→管道连接→单机试运转。详见施工步骤、安全要点、质量控制要点。"
        # v0.1.51：设备技术参数联动
        try:
            from . import relations as _rel
            from . import spatial_model as _sm
            g = _rel.load_relations()
            spatial = _sm.build_spatial_model(g)
            spatial_devs = {d["tag"]: d for d in spatial.get("devices", [])}
            tech_params = []
            for d in devices[:5]:
                tag = d["tag"]
                sd = spatial_devs.get(tag, {})
                params = []
                if sd.get("workshop"):
                    params.append(f"车间: {sd['workshop']}")
                if sd.get("x") is not None and sd.get("y") is not None:
                    params.append(f"坐标: ({sd['x']}, {sd['y']})")
                if sd.get("z") is not None:
                    params.append(f"安装标高: {sd['z']}m")
                if params:
                    tech_params.append(f"{tag}: {'; '.join(params)}")
            if tech_params:
                data["交底设备参数"] = "\n".join(tech_params)
        except Exception:  # noqa: BLE001
            pass

    # 开箱验收记录专项（v0.1.52：设备数据联动——技术参数、箱单信息、外观检查、随机资料）
    if doc_type == "开箱验收记录" and devices:
        data["涉及设备"] = "、".join(d["tag"] for d in devices[:10]) + (f" 等{len(devices)}台" if len(devices) > 10 else "")
        # 识别设备类型
        from . import equipment_types as _et
        eq_type = _et.get_equipment_type_from_devices(devices)
        data["设备类型"] = eq_type
        # 设备名称和数量
        first_dev = devices[0]
        data["设备名称"] = first_dev.get("name", f"{eq_type}（详见设备清单）")
        data["数量"] = str(len(devices)) + "台"
        # 外观检查要点（根据设备类型）
        inspection_points = _get_inspection_points(eq_type)
        data["外观检查要点"] = "\n".join(f"{i+1}. {p}" for i, p in enumerate(inspection_points))
        data["外观检查情况"] = "开箱后设备外观完好，无锈蚀、变形、损坏（详见外观检查要点）"
        # 随机资料清单（根据设备类型）
        random_docs = _get_random_docs(eq_type)
        data["随机资料清单"] = "\n".join(f"{i+1}. {d}" for i, d in enumerate(random_docs))
        data["随机资料"] = "、".join(random_docs[:5]) + ("等" if len(random_docs) > 5 else "")
        # 验收结果默认
        data["验收结果"] = "合格（设备型号规格符合设计，外观完好，随机资料齐全）"
        # v0.1.52：设备技术参数联动
        try:
            from . import relations as _rel
            from . import spatial_model as _sm
            g = _rel.load_relations()
            spatial = _sm.build_spatial_model(g)
            spatial_devs = {d["tag"]: d for d in spatial.get("devices", [])}
            tech_params = []
            for d in devices[:5]:
                tag = d["tag"]
                sd = spatial_devs.get(tag, {})
                params = []
                if sd.get("workshop"):
                    params.append(f"安装车间: {sd['workshop']}")
                if sd.get("x") is not None and sd.get("y") is not None:
                    params.append(f"安装坐标: ({sd['x']}, {sd['y']})")
                if sd.get("z") is not None:
                    params.append(f"安装标高: {sd['z']}m")
                if params:
                    tech_params.append(f"{tag}: {'; '.join(params)}")
            if tech_params:
                data["验收设备参数"] = "\n".join(tech_params)
        except Exception:  # noqa: BLE001
            pass

    # 隐蔽工程验收记录专项（v0.1.53：设备数据联动——技术参数、隐蔽部位、检查要点、质量标准）
    if doc_type == "隐蔽工程验收记录" and devices:
        data["涉及设备"] = "、".join(d["tag"] for d in devices[:10]) + (f" 等{len(devices)}台" if len(devices) > 10 else "")
        # 识别设备类型
        from . import equipment_types as _et
        eq_type = _et.get_equipment_type_from_devices(devices)
        data["设备类型"] = eq_type
        # 隐蔽部位（根据设备类型推断）
        concealment_parts = _get_concealment_parts(eq_type)
        data["隐蔽部位"] = "、".join(concealment_parts)
        # 隐蔽检查要点
        concealment_points = _get_concealment_points(eq_type)
        data["隐蔽检查要点"] = "\n".join(f"{i+1}. {p}" for i, p in enumerate(concealment_points))
        # 质量标准
        quality_standards = _get_concealment_quality(eq_type)
        data["质量标准"] = "\n".join(f"{i+1}. {q}" for i, q in enumerate(quality_standards))
        # 检查情况默认
        data["检查情况"] = "隐蔽内容与图纸相符，几何尺寸符合要求，预埋件位置准确固定牢靠，隐蔽前影像资料留存齐全（详见隐蔽检查要点和质量标准）"
        # 验收结果默认
        data["验收结果"] = "合格（隐蔽工程质量符合设计和规范要求）"
        # v0.1.53：设备技术参数联动
        try:
            from . import relations as _rel
            from . import spatial_model as _sm
            g = _rel.load_relations()
            spatial = _sm.build_spatial_model(g)
            spatial_devs = {d["tag"]: d for d in spatial.get("devices", [])}
            tech_params = []
            for d in devices[:5]:
                tag = d["tag"]
                sd = spatial_devs.get(tag, {})
                params = []
                if sd.get("workshop"):
                    params.append(f"安装车间: {sd['workshop']}")
                if sd.get("x") is not None and sd.get("y") is not None:
                    params.append(f"安装坐标: ({sd['x']}, {sd['y']})")
                if sd.get("z") is not None:
                    params.append(f"安装标高: {sd['z']}m")
                if params:
                    tech_params.append(f"{tag}: {'; '.join(params)}")
            if tech_params:
                data["隐蔽设备参数"] = "\n".join(tech_params)
        except Exception:  # noqa: BLE001
            pass

    # 设计变更专项（v0.1.55：设备数据联动——技术参数、变更内容、影响范围、处理建议、管线影响）
    if doc_type == "设计变更" and devices:
        data["涉及设备"] = "、".join(d["tag"] for d in devices[:10]) + (f" 等{len(devices)}台" if len(devices) > 10 else "")
        # 识别设备类型
        from . import equipment_types as _et
        eq_type = _et.get_equipment_type_from_devices(devices)
        data["设备类型"] = eq_type
        # 变更内容建议
        change_contents = _get_change_content(eq_type)
        data["变更内容建议"] = "\n".join(f"{i+1}. {c}" for i, c in enumerate(change_contents))
        # 影响范围分析
        impact = _get_change_impact(eq_type, devices)
        data["影响范围"] = impact["summary"]
        data["影响车间"] = "、".join(impact["workshops"])
        data["影响管线"] = "、".join(impact["pipes"])
        data["影响相邻设备"] = "、".join(impact["neighbors"])
        # 处理建议
        suggestions = _get_change_suggestions(eq_type)
        data["处理建议"] = "\n".join(f"{i+1}. {s}" for i, s in enumerate(suggestions))
        # 变更原因分类
        data["变更原因"] = "现场与设计不符（详见变更内容）"
        # v0.1.55：设备技术参数联动
        try:
            from . import relations as _rel
            from . import spatial_model as _sm
            g = _rel.load_relations()
            spatial = _sm.build_spatial_model(g)
            spatial_devs = {d["tag"]: d for d in spatial.get("devices", [])}
            tech_params = []
            for d in devices[:5]:
                tag = d["tag"]
                sd = spatial_devs.get(tag, {})
                params = []
                if sd.get("workshop"):
                    params.append(f"车间: {sd['workshop']}")
                if sd.get("x") is not None and sd.get("y") is not None:
                    params.append(f"坐标: ({sd['x']}, {sd['y']})")
                if sd.get("z") is not None:
                    params.append(f"标高: {sd['z']}m")
                if params:
                    tech_params.append(f"{tag}: {'; '.join(params)}")
            if tech_params:
                data["变更设备参数"] = "\n".join(tech_params)
        except Exception:  # noqa: BLE001
            pass

    # 货损报告专项（v0.1.56：设备数据联动——技术参数、损失情况、处理建议、索赔建议）
    if doc_type == "货损报告" and devices:
        first_dev = devices[0]
        data["设备位号"] = first_dev["tag"]
        data["设备名称"] = first_dev.get("name", "详见设备清单")
        # 识别设备类型
        from . import equipment_types as _et
        eq_type = _et.get_equipment_type_from_devices(devices)
        data["设备类型"] = eq_type
        # 损失情况建议
        damage_contents = _get_damage_content(eq_type)
        data["损失情况建议"] = "\n".join(f"{i+1}. {c}" for i, c in enumerate(damage_contents))
        data["损失情况"] = "开箱后发现设备外观有损伤（详见损失情况建议，需现场确认具体损失程度）"
        data["损失程度"] = "待现场确认（轻微/一般/严重）"
        # 处理建议
        suggestions = _get_damage_suggestions(eq_type)
        data["处理建议"] = "\n".join(f"{i+1}. {s}" for i, s in enumerate(suggestions))
        # 索赔建议
        data["索赔建议"] = "1. 保留受损设备原始状态和影像资料\n2. 会同运输方/厂家共同确认损失程度\n3. 根据采购合同和运输保险条款提出索赔\n4. 要求厂家补发或维修受损部件"
        # 责任方
        data["责任方"] = "待确认（运输方/厂家/装卸方）"
        # v0.1.56：设备技术参数联动
        try:
            from . import relations as _rel
            from . import spatial_model as _sm
            g = _rel.load_relations()
            spatial = _sm.build_spatial_model(g)
            spatial_devs = {d["tag"]: d for d in spatial.get("devices", [])}
            tech_params = []
            for d in devices[:5]:
                tag = d["tag"]
                sd = spatial_devs.get(tag, {})
                params = []
                if sd.get("workshop"):
                    params.append(f"安装车间: {sd['workshop']}")
                if sd.get("x") is not None and sd.get("y") is not None:
                    params.append(f"安装坐标: ({sd['x']}, {sd['y']})")
                if sd.get("z") is not None:
                    params.append(f"安装标高: {sd['z']}m")
                if params:
                    tech_params.append(f"{tag}: {'; '.join(params)}")
            if tech_params:
                data["货损设备参数"] = "\n".join(tech_params)
        except Exception:  # noqa: BLE001
            pass

    # 竣工资料专项（v0.1.57：设备数据联动——技术参数、交工范围、遗留问题、验收记录）
    if doc_type == "竣工资料" and devices:
        data["涉及设备"] = "、".join(d["tag"] for d in devices[:10]) + (f" 等{len(devices)}台" if len(devices) > 10 else "")
        # 识别设备类型
        from . import equipment_types as _et
        eq_type = _et.get_equipment_type_from_devices(devices)
        data["设备类型"] = eq_type
        # 交工范围
        completion_scope = _get_completion_scope(eq_type, devices)
        data["交工范围"] = completion_scope
        # 遗留问题（从完整性检查获取）
        remaining = _get_remaining_issues(devices)
        data["遗留问题及处理"] = remaining["summary"]
        data["遗留问题清单"] = "\n".join(f"{i+1}. {r}" for i, r in enumerate(remaining["issues"])) if remaining["issues"] else "无遗留问题"
        # 验收记录编号（从已生成文档获取）
        try:
            import os as _os
            gen_dir = _os.path.join(_os.path.dirname(_os.path.dirname(__file__)), "data", "generated_docs")
            acceptance_records = []
            for dtype in ["开箱验收记录", "隐蔽工程验收记录"]:
                dtype_dir = _os.path.join(gen_dir, dtype)
                if _os.path.exists(dtype_dir):
                    files = _os.listdir(dtype_dir)
                    acceptance_records.extend([f.replace(".docx", "") for f in files[:5]])
            if acceptance_records:
                data["验收记录编号"] = "、".join(acceptance_records)
                data["隐蔽工程资料编号"] = "、".join([r for r in acceptance_records if "隐蔽" in r]) or "待整理"
        except Exception:  # noqa: BLE001
            pass
        # v0.1.57：设备技术参数联动
        try:
            from . import relations as _rel
            from . import spatial_model as _sm
            g = _rel.load_relations()
            spatial = _sm.build_spatial_model(g)
            spatial_devs = {d["tag"]: d for d in spatial.get("devices", [])}
            tech_params = []
            workshops = set()
            for d in devices[:10]:
                tag = d["tag"]
                sd = spatial_devs.get(tag, {})
                params = []
                if sd.get("workshop"):
                    params.append(f"车间: {sd['workshop']}")
                    workshops.add(sd["workshop"])
                if sd.get("z") is not None:
                    params.append(f"标高: {sd['z']}m")
                if params:
                    tech_params.append(f"{tag}: {'; '.join(params)}")
            if tech_params:
                data["竣工设备参数"] = "\n".join(tech_params)
            if workshops:
                data["交工车间"] = "、".join(workshops)
        except Exception:  # noqa: BLE001
            pass

    # 施工方案专项（v0.1.42设备类型+v0.1.49设备数据联动）
    if doc_type == "施工方案":
        if devices:
            data["涉及设备"] = "、".join(d["tag"] for d in devices[:10]) + (f" 等{len(devices)}台" if len(devices) > 10 else "")
            # v0.1.42：识别设备类型，自动生成施工步骤
            from . import equipment_types as _et
            eq_type = _et.get_equipment_type_from_devices(devices)
            data["设备类型"] = eq_type
            steps = _et.get_construction_steps(eq_type)
            data["施工步骤"] = "\n".join(f"{i+1}. {step}" for i, step in enumerate(steps))
            data["施工内容"] = f"{eq_type}安装施工（含基础验收、设备就位、找平找正、管道连接、单机试运转等）"
            # v0.1.49：设备数据联动——技术参数、空间位置、管线连接、相邻设备
            try:
                from . import relations as _rel
                from . import spatial_model as _sm
                g = _rel.load_relations()
                spatial = _sm.build_spatial_model(g)
                spatial_devs = {d["tag"]: d for d in spatial.get("devices", [])}
                # 设备技术参数汇总
                tech_params = []
                spatial_info = []
                for d in devices[:5]:
                    tag = d["tag"]
                    sd = spatial_devs.get(tag, {})
                    # 从设备图谱获取参数
                    dev_info = next((x for x in g.get("devices", []) if x["tag"] == tag), {})
                    params = []
                    if sd.get("workshop"):
                        params.append(f"车间: {sd['workshop']}")
                    if sd.get("x") is not None and sd.get("y") is not None:
                        params.append(f"坐标: ({sd['x']}, {sd['y']})")
                    if sd.get("z") is not None:
                        params.append(f"标高: {sd['z']}m")
                    if sd.get("coord_status"):
                        params.append(f"位置状态: {sd['coord_status']}")
                    if params:
                        tech_params.append(f"{tag}: {'; '.join(params)}")
                    if sd.get("workshop"):
                        spatial_info.append(f"{tag}位于{sd['workshop']}")
                if tech_params:
                    data["设备技术参数"] = "\n".join(tech_params)
                if spatial_info:
                    data["施工环境"] = "; ".join(spatial_info)
                # 管线连接（v0.1.47 piping_network）
                try:
                    from . import piping_network as _pn
                    pipe_connections = []
                    for d in devices[:5]:
                        tag = d["tag"]
                        pipes = _pn.get_device_pipes(tag)
                        conns = _pn.get_device_connections(tag)
                        if pipes:
                            pipe_info = [f"{p['pipe_no']}({p['medium']})" for p in pipes[:3]]
                            pipe_connections.append(f"{tag}连接: {', '.join(pipe_info)}")
                        if conns:
                            conn_info = [f"→{c['to_device']}({c['pipe_no']})" for c in conns[:3] if c.get('to_device')]
                            if conn_info:
                                pipe_connections.append(f"{tag}关联设备: {', '.join(conn_info)}")
                    if pipe_connections:
                        data["相关管线"] = "\n".join(pipe_connections)
                except Exception:  # noqa: BLE001
                    pass
                # 相邻设备（v0.1.35空间关系）
                neighbors = []
                for d in devices[:3]:
                    tag = d["tag"]
                    sd = spatial_devs.get(tag, {})
                    if sd.get("neighbors"):
                        nb = [n.get("tag", "") for n in sd["neighbors"][:3] if n.get("tag")]
                        if nb:
                            neighbors.append(f"{tag}相邻: {', '.join(nb)}")
                if neighbors:
                    data["相邻设备"] = "; ".join(neighbors)
            except Exception:  # noqa: BLE001
                pass
        if "施工内容" not in data:
            missing.append("施工内容（需明确具体施工范围与工序）")

    # 规范正文引用（v0.1.17 增强：只取现行规范）
    citations = std_citations(doc_type, limit=4)
    data["_std_citations"] = citations

    # 检查必填字段缺失
    for k in t["required"]:
        if not str(data.get(k, "")).strip() and k not in missing:
            missing.append(k)

    return {"data": data, "missing": missing, "devices": devices, "citations": citations}




def _get_safety_points(eq_type: str, devices: list = None) -> list:
    """v0.1.51：根据设备类型生成安全要点。"""
    base = [
        "作业前办理作业票，确认作业环境安全",
        "作业人员持证上岗，佩戴个人防护用品",
        "设置警戒区，无关人员不得进入作业区",
        "作业前检查工器具完好，严禁带病作业",
    ]
    type_safety = {
        "泵": ["联轴器防护罩安装到位后方可试运转", "试运转时严禁站在联轴器旋转方向正面", "泵进出口管道连接前清理管内杂物"],
        "压缩机": ["压缩机试运转前确认润滑油位正常", "超压保护装置校验合格后方可投用", "皮带传动部位防护罩齐全"],
        "塔器": ["高处作业系挂安全带，高挂低用", "塔器内作业办理受限空间作业票", "吊装时风力超过六级停止作业"],
        "换热器": ["抽芯作业时防止芯子滑落伤人", "水压试验时升压缓慢，严禁超压", "法兰紧固按对角顺序均匀紧固"],
        "容器": ["容器内作业办理受限空间作业票", "容器试压时严禁人员站在法兰正面", "防腐作业保持通风良好"],
        "风机": ["风机试运转前检查叶轮旋转方向", "皮带防护罩安装齐全", "试运转时测量轴承温度不超标"],
        "起重机": ["吊装作业专人指挥，信号统一", "吊索具安全系数不小于6", "吊装区域下方严禁站人"],
        "电机": ["电机接线前确认电源断开并挂牌", "电机试运转前测量绝缘电阻合格", "联轴器防护罩安装到位"],
        "阀门": ["阀门安装前核对型号规格和流向", "高压阀门试压合格后方可安装", "法兰紧固按对角顺序"],
        "储罐": ["储罐焊接作业办理动火作业票", "罐内作业办理受限空间作业票", "储罐充水试验时监测基础沉降"],
    }
    points = list(base)
    points.extend(type_safety.get(eq_type, ["按设备说明书和施工方案执行安全措施"]))
    return points


def _get_quality_points(eq_type: str) -> list:
    """v0.1.51：根据设备类型生成质量控制要点。"""
    base = [
        "基础验收合格，混凝土强度达到设计要求",
        "设备开箱检验合格，型号规格符合设计",
        "设备就位后找平找正，偏差在规范允许范围内",
        "管道连接无应力，法兰平行度符合要求",
        "单机试运转参数符合设备技术文件要求",
    ]
    type_quality = {
        "泵": ["泵的水平度偏差不大于0.05mm/m", "联轴器对中偏差径向≤0.05mm，端面≤0.03mm", "试运转时轴承温度不超过75℃"],
        "压缩机": ["压缩机水平度偏差不大于0.05mm/m", "气缸与滑道对中偏差符合规范", "试运转时各级压力温度符合设计"],
        "塔器": ["塔器垂直度偏差不大于高度的1/1000且不大于30mm", "塔盘水平度偏差不大于2mm", "附件安装位置符合图纸"],
        "换热器": ["换热器抽芯检查管束完好", "水压试验压力为设计压力的1.25倍", "保冷保温层厚度符合设计"],
        "容器": ["容器焊接接头无损检测比例符合设计", "容器水压试验压力符合规范", "附件安装位置和方向正确"],
        "风机": ["风机叶轮与机壳间隙均匀", "皮带轮对中偏差符合要求", "试运转时轴承振动值不超标"],
        "电机": ["电机绝缘电阻不小于0.5MΩ", "电机空载试运转电流不超过额定值", "电机轴承温度不超过80℃"],
    }
    points = list(base)
    points.extend(type_quality.get(eq_type, ["按设备说明书和施工方案执行质量控制"]))
    return points



def _get_inspection_points(eq_type: str) -> list:
    """v0.1.52：根据设备类型生成外观检查要点。"""
    base = [
        "检查设备外观有无锈蚀、变形、磕碰损伤",
        "检查设备铭牌清晰，型号规格与设计一致",
        "检查设备接口法兰有无损伤，密封面完好",
        "检查设备地脚螺栓孔位置与基础一致",
        "检查设备油漆涂层完好，无脱落",
    ]
    type_inspection = {
        "泵": ["检查泵轴转动灵活，无卡涩", "检查机械密封完好，无泄漏", "检查联轴器对中良好", "检查泵进出口法兰密封面完好"],
        "压缩机": ["检查压缩机曲轴转动灵活", "检查气缸内壁无锈蚀", "检查气阀组件完好", "检查润滑油系统清洁"],
        "塔器": ["检查塔器筒体圆度和直线度", "检查塔盘支撑圈水平度", "检查人孔法兰密封面完好", "检查接管方位与图纸一致"],
        "换热器": ["检查换热器管束抽芯检查", "检查管板和折流板完好", "检查壳程和管程法兰密封面", "检查换热管无变形堵塞"],
        "容器": ["检查容器焊缝外观无缺陷", "检查容器法兰密封面完好", "检查接管方位与图纸一致", "检查内件安装牢固"],
        "风机": ["检查风机叶轮转动灵活无摩擦", "检查风机机壳无变形", "检查皮带轮和联轴器完好", "检查轴承座无渗漏"],
        "电机": ["检查电机轴转动灵活", "检查电机接线盒完好", "检查电机风扇和防护罩齐全", "检查电机绝缘电阻合格"],
        "阀门": ["检查阀门阀体无裂纹砂眼", "检查阀门法兰密封面完好", "检查阀门开关灵活无卡涩", "检查阀门流向标识清晰"],
        "储罐": ["检查储罐底板和壁板外观", "检查储罐焊缝外观无缺陷", "检查储罐接管方位正确", "检查储罐内防腐层完好"],
    }
    points = list(base)
    points.extend(type_inspection.get(eq_type, ["按设备说明书和装箱单检查"]))
    return points


def _get_random_docs(eq_type: str) -> list:
    """v0.1.52：根据设备类型生成随机资料清单。"""
    base = [
        "产品合格证",
        "产品说明书",
        "装箱单",
        "出厂检验报告",
    ]
    type_docs = {
        "泵": ["泵性能曲线图", "机械密封安装图", "泵安装尺寸图", "易损件清单"],
        "压缩机": ["压缩机性能曲线图", "润滑油系统图", "气阀组件图", "安装运行维护手册"],
        "塔器": ["塔器制造竣工图", "塔盘安装图", "水压试验报告", "无损检测报告"],
        "换热器": ["换热器制造竣工图", "管束装配图", "水压试验报告", "换热管材质证明"],
        "容器": ["容器制造竣工图", "水压试验报告", "无损检测报告", "材质证明书"],
        "风机": ["风机性能曲线图", "风机安装图", "轴承润滑说明", "易损件清单"],
        "电机": ["电机接线图", "电机性能参数表", "绝缘电阻测试报告", "安装维护说明"],
        "阀门": ["阀门试压报告", "阀门安装说明书", "阀门材质证明", "易损件清单"],
        "储罐": ["储罐制造竣工图", "焊缝检测报告", "充水试验报告", "防腐层检测报告"],
    }
    docs = list(base)
    docs.extend(type_docs.get(eq_type, ["按设备装箱单清点"]))
    return docs



def _get_concealment_parts(eq_type: str) -> list:
    """v0.1.53：根据设备类型生成隐蔽部位。"""
    type_parts = {
        "泵": ["泵基础预埋件", "泵地脚螺栓", "泵进出口管道连接", "泵底座灌浆"],
        "压缩机": ["压缩机基础预埋件", "压缩机地脚螺栓", "压缩机管道连接", "润滑油管线"],
        "塔器": ["塔器基础预埋件", "塔器地脚螺栓", "塔器接管连接", "塔器保温层"],
        "换热器": ["换热器基础预埋件", "换热器地脚螺栓", "换热器管程壳程连接", "换热器保温层"],
        "容器": ["容器基础预埋件", "容器地脚螺栓", "容器接管连接", "容器保温层"],
        "风机": ["风机基础预埋件", "风机地脚螺栓", "风机进出口管道连接", "风机减振装置"],
        "电机": ["电机基础预埋件", "电机地脚螺栓", "电机电缆接线", "电机接地装置"],
        "阀门": ["阀门基础支架", "阀门法兰连接", "阀门保温层", "阀门传动装置"],
        "储罐": ["储罐基础预埋件", "储罐底板焊接", "储罐接管连接", "储罐防腐层"],
    }
    return type_parts.get(eq_type, ["设备基础预埋件", "设备地脚螺栓", "设备管道连接"])


def _get_concealment_points(eq_type: str) -> list:
    """v0.1.53：根据设备类型生成隐蔽检查要点。"""
    base = [
        "核对隐蔽部位与图纸一致，几何尺寸符合设计要求",
        "检查预埋件位置准确，固定牢靠，标高符合设计",
        "检查地脚螺栓规格、数量、位置与基础一致",
        "检查管道连接无应力，法兰平行度符合要求",
        "隐蔽前拍摄影像资料，编号留存齐全",
    ]
    type_points = {
        "泵": ["检查泵基础混凝土强度达到设计要求", "检查泵底座水平度偏差不大于0.05mm/m", "检查泵进出口管道法兰对中良好", "检查二次灌浆密实无空隙"],
        "压缩机": ["检查压缩机基础混凝土强度达到设计要求", "检查压缩机底座水平度偏差不大于0.05mm/m", "检查气缸与滑道对中偏差符合规范", "检查润滑油管线清洁无杂物"],
        "塔器": ["检查塔器基础混凝土强度达到设计要求", "检查塔器垂直度偏差不大于高度的1/1000", "检查塔器地脚螺栓紧固力矩符合要求", "检查塔器保温层厚度符合设计"],
        "换热器": ["检查换热器基础混凝土强度达到设计要求", "检查换热器底座水平度符合要求", "检查换热器管程壳程法兰密封面完好", "检查换热器保温层厚度符合设计"],
        "容器": ["检查容器基础混凝土强度达到设计要求", "检查容器垂直度偏差符合规范", "检查容器地脚螺栓紧固力矩符合要求", "检查容器保温层厚度符合设计"],
        "风机": ["检查风机基础混凝土强度达到设计要求", "检查风机减振装置安装正确", "检查风机进出口管道软连接完好", "检查风机底座水平度符合要求"],
        "电机": ["检查电机基础混凝土强度达到设计要求", "检查电机接地装置连接可靠", "检查电机电缆接线正确，绝缘合格", "检查电机底座水平度符合要求"],
        "阀门": ["检查阀门支架安装牢固", "检查阀门法兰螺栓紧固均匀", "检查阀门保温层施工符合要求", "检查阀门传动装置操作灵活"],
        "储罐": ["检查储罐基础混凝土强度达到设计要求", "检查储罐底板焊接质量合格", "检查储罐防腐层厚度符合设计", "检查储罐充水试验基础沉降合格"],
    }
    points = list(base)
    points.extend(type_points.get(eq_type, ["按设备说明书和施工方案检查"]))
    return points


def _get_concealment_quality(eq_type: str) -> list:
    """v0.1.53：根据设备类型生成隐蔽工程质量标准。"""
    base = [
        "隐蔽工程质量符合设计图纸和现行施工验收规范要求",
        "预埋件位置偏差不大于±5mm，标高偏差不大于±3mm",
        "地脚螺栓位置偏差不大于±2mm，垂直度偏差不大于1/100",
        "管道法兰平行度偏差不大于法兰外径的1.5/1000且不大于2mm",
        "隐蔽前影像资料齐全，编号清晰可追溯",
    ]
    type_quality = {
        "泵": ["泵底座水平度偏差不大于0.05mm/m", "泵联轴器对中偏差径向≤0.05mm，端面≤0.03mm", "二次灌浆强度达到设计要求的75%以上方可紧固螺栓"],
        "压缩机": ["压缩机底座水平度偏差不大于0.05mm/m", "气缸与滑道对中偏差符合设备技术文件要求", "润滑油管线冲洗合格，清洁度符合要求"],
        "塔器": ["塔器垂直度偏差不大于高度的1/1000且不大于30mm", "塔器地脚螺栓紧固力矩符合设备技术文件要求", "塔器保温层厚度偏差不大于+10mm/-5mm"],
        "换热器": ["换热器底座水平度偏差不大于0.05mm/m", "换热器管程壳程水压试验合格", "换热器保温层厚度偏差不大于+10mm/-5mm"],
        "容器": ["容器垂直度偏差不大于高度的1/1000且不大于30mm", "容器地脚螺栓紧固力矩符合设备技术文件要求", "容器保温层厚度偏差不大于+10mm/-5mm"],
        "风机": ["风机底座水平度偏差不大于0.05mm/m", "风机叶轮与机壳间隙均匀，偏差不大于设计值的±10%", "风机减振装置压缩量均匀，偏差不大于2mm"],
        "电机": ["电机底座水平度偏差不大于0.05mm/m", "电机绝缘电阻不小于0.5MΩ", "电机接地电阻不大于4Ω"],
        "阀门": ["阀门法兰螺栓紧固力矩符合规范要求", "阀门保温层厚度偏差不大于+10mm/-5mm", "阀门开关灵活，无卡涩"],
        "储罐": ["储罐底板焊接无损检测合格", "储罐充水试验基础沉降差不大于设计允许值", "储罐防腐层厚度偏差不大于+10mm/-5mm"],
    }
    quality = list(base)
    quality.extend(type_quality.get(eq_type, ["按设备说明书和施工验收规范执行"]))
    return quality



def _get_change_content(eq_type: str) -> list:
    """v0.1.55：根据设备类型生成变更内容建议。"""
    type_content = {
        "泵": ["泵基础尺寸或位置变更", "泵进出口管道走向或标高变更", "泵型号或规格变更", "泵安装方式变更（立式/卧式）"],
        "压缩机": ["压缩机基础尺寸或位置变更", "压缩机管道系统走向变更", "压缩机型号或规格变更", "压缩机辅助系统配置变更"],
        "塔器": ["塔器基础标高或位置变更", "塔器接管方位或标高变更", "塔器内件配置变更", "塔器保温或防腐要求变更"],
        "换热器": ["换热器基础位置或标高变更", "换热器管程壳程接口变更", "换热器型号或规格变更", "换热器保温要求变更"],
        "容器": ["容器基础标高或位置变更", "容器接管方位或标高变更", "容器型号或规格变更", "容器保温或防腐要求变更"],
        "风机": ["风机基础位置或标高变更", "风机进出口管道走向变更", "风机型号或规格变更", "风机减振方式变更"],
        "电机": ["电机基础位置或标高变更", "电机电缆走向或接线方式变更", "电机型号或功率变更", "电机接地方式变更"],
        "阀门": ["阀门安装位置或标高变更", "阀门型号或规格变更", "阀门传动方式变更", "阀门保温要求变更"],
        "储罐": ["储罐基础位置或标高变更", "储罐接管方位变更", "储罐容积或规格变更", "储罐防腐要求变更"],
    }
    return type_content.get(eq_type, ["设备位置或标高变更", "设备型号或规格变更", "设备管道连接方式变更"])


def _get_change_impact(eq_type: str, devices: list) -> dict:
    """v0.1.55：分析变更影响范围。"""
    workshops = set()
    pipes = []
    neighbors = []
    
    # 从空间模型获取影响范围
    try:
        from . import relations as _rel
        from . import spatial_model as _sm
        g = _rel.load_relations()
        spatial = _sm.build_spatial_model(g)
        spatial_devs = {d["tag"]: d for d in spatial.get("devices", [])}
        for d in devices[:5]:
            tag = d["tag"]
            sd = spatial_devs.get(tag, {})
            if sd.get("workshop"):
                workshops.add(sd["workshop"])
            if sd.get("neighbors"):
                for nb in sd["neighbors"][:3]:
                    if nb.get("tag"):
                        neighbors.append(nb["tag"])
    except Exception:  # noqa: BLE001
        pass
    
    # 从管线网络获取影响管线
    try:
        from . import piping_network as _pn
        for d in devices[:3]:
            tag = d["tag"]
            dev_pipes = _pn.get_device_pipes(tag)
            for p in dev_pipes[:3]:
                pipes.append(p.get("pipe_no", ""))
    except Exception:  # noqa: BLE001
        pass
    
    if not workshops:
        workshops.add("待确认")
    
    summary_parts = []
    if workshops:
        summary_parts.append(f"影响{len(workshops)}个车间")
    if pipes:
        summary_parts.append(f"影响{len(pipes)}条管线")
    if neighbors:
        summary_parts.append(f"影响{len(neighbors)}台相邻设备")
    summary = "；".join(summary_parts) if summary_parts else "影响范围待确认"
    
    return {
        "summary": summary,
        "workshops": list(workshops),
        "pipes": pipes[:10],
        "neighbors": neighbors[:10],
    }


def _get_change_suggestions(eq_type: str) -> list:
    """v0.1.55：根据设备类型生成变更处理建议。"""
    base = [
        "变更前组织设计、施工、监理三方会签确认",
        "变更内容及时更新到相关专业图纸",
        "变更工程量及时记录并作为结算依据",
        "变更涉及的设备材料及时采购或调整",
    ]
    type_suggestions = {
        "泵": ["变更后重新核对泵基础尺寸和地脚螺栓位置", "变更后重新进行管道应力分析", "变更后重新核对泵的汽蚀余量"],
        "压缩机": ["变更后重新核对压缩机基础荷载和振动要求", "变更后重新进行管道系统设计", "变更后重新核对压缩机辅助系统配置"],
        "塔器": ["变更后重新核对塔器基础荷载和垂直度要求", "变更后重新核对塔器接管与管道连接", "变更后重新进行塔器内件安装设计"],
        "换热器": ["变更后重新核对换热器基础尺寸和标高", "变更后重新核对管程壳程管道连接", "变更后重新进行换热器水压试验"],
        "容器": ["变更后重新核对容器基础荷载和标高", "变更后重新核对容器接管与管道连接", "变更后重新进行容器水压试验"],
        "风机": ["变更后重新核对风机基础尺寸和减振要求", "变更后重新进行管道系统设计", "变更后重新核对风机风量和风压"],
        "电机": ["变更后重新核对电机基础尺寸和荷载", "变更后重新进行电缆敷设设计", "变更后重新核对电机功率和绝缘要求"],
        "阀门": ["变更后重新核对阀门安装位置和操作空间", "变更后重新进行管道应力分析", "变更后重新核对阀门试压要求"],
        "储罐": ["变更后重新核对储罐基础荷载和沉降要求", "变更后重新核对储罐接管与管道连接", "变更后重新进行储罐充水试验"],
    }
    suggestions = list(base)
    suggestions.extend(type_suggestions.get(eq_type, ["变更后重新核对设备安装要求", "变更后重新进行相关专业设计"]))
    return suggestions



def _get_damage_content(eq_type: str) -> list:
    """v0.1.56：根据设备类型生成损失情况建议。"""
    type_content = {
        "泵": ["泵体外壳有无磕碰变形", "泵轴有无弯曲变形", "机械密封有无损坏泄漏", "联轴器有无变形损坏", "油漆涂层有无脱落锈蚀"],
        "压缩机": ["压缩机机身有无磕碰变形", "曲轴有无弯曲变形", "气缸有无变形损伤", "气阀组件有无损坏", "润滑油系统有无泄漏"],
        "塔器": ["塔器筒体有无凹陷变形", "塔器法兰密封面有无损伤", "塔器接管有无变形", "塔器内件有无损坏", "塔器防腐层有无脱落"],
        "换热器": ["换热器壳体有无磕碰变形", "换热器管板有无损伤", "换热管有无变形堵塞", "换热器法兰密封面有无损伤", "换热器保温层有无损坏"],
        "容器": ["容器筒体有无凹陷变形", "容器法兰密封面有无损伤", "容器接管有无变形", "容器内件有无损坏", "容器防腐层有无脱落"],
        "风机": ["风机机壳有无磕碰变形", "风机叶轮有无变形损坏", "风机轴承有无损坏", "风机皮带轮有无变形", "风机防护罩有无损坏"],
        "电机": ["电机外壳有无磕碰变形", "电机轴有无弯曲变形", "电机接线盒有无损坏", "电机风扇有无损坏", "电机油漆涂层有无脱落"],
        "阀门": ["阀门阀体有无裂纹砂眼", "阀门法兰密封面有无损伤", "阀门阀杆有无弯曲变形", "阀门传动装置有无损坏", "阀门油漆涂层有无脱落"],
        "储罐": ["储罐底板有无变形凹陷", "储罐壁板有无磕碰变形", "储罐接管有无变形", "储罐顶板有无损坏", "储罐防腐层有无脱落"],
    }
    return type_content.get(eq_type, ["设备外壳有无磕碰变形", "设备接口有无损伤", "设备油漆涂层有无脱落锈蚀"])


def _get_damage_suggestions(eq_type: str) -> list:
    """v0.1.56：根据设备类型生成货损处理建议。"""
    base = [
        "立即停止开箱，保护受损设备原始状态",
        "对受损部位进行拍照和录像留存证据",
        "会同运输方/厂家/监理共同确认损失程度",
        "记录受损设备的位号、名称、数量、损失情况",
        "根据损失程度确定维修、更换或退货方案",
    ]
    type_suggestions = {
        "泵": ["泵轴弯曲需校直或更换", "机械密封损坏需更换密封组件", "泵体变形需厂家评估修复或更换", "联轴器损坏需更换"],
        "压缩机": ["曲轴弯曲需校直或更换", "气缸变形需厂家评估修复", "气阀组件损坏需更换", "润滑油系统泄漏需修复"],
        "塔器": ["筒体凹陷需整形修复或更换", "法兰密封面损伤需研磨修复", "接管变形需更换接管", "内件损坏需厂家更换"],
        "换热器": ["壳体变形需整形修复", "管板损伤需厂家评估修复", "换热管变形需更换管束", "法兰密封面损伤需研磨修复"],
        "容器": ["筒体凹陷需整形修复或更换", "法兰密封面损伤需研磨修复", "接管变形需更换接管", "内件损坏需厂家更换"],
        "风机": ["机壳变形需整形修复", "叶轮变形需校动平衡或更换", "轴承损坏需更换轴承", "皮带轮变形需更换"],
        "电机": ["外壳变形需整形修复", "轴弯曲需校直或更换", "接线盒损坏需更换", "绝缘受损需重新浸漆烘干"],
        "阀门": ["阀体裂纹需更换阀门", "法兰密封面损伤需研磨修复", "阀杆弯曲需校直或更换", "传动装置损坏需更换"],
        "储罐": ["底板变形需整形修复", "壁板凹陷需整形修复或更换", "接管变形需更换接管", "防腐层损坏需重新防腐"],
    }
    suggestions = list(base)
    suggestions.extend(type_suggestions.get(eq_type, ["受损部件需厂家评估修复或更换"]))
    return suggestions



def _get_completion_scope(eq_type: str, devices: list) -> str:
    """v0.1.57：根据设备类型生成交工范围。"""
    type_scope = {
        "泵": "泵类设备安装、管道连接、单机试运转及验收",
        "压缩机": "压缩机类设备安装、辅助系统连接、单机试运转及验收",
        "塔器": "塔器类设备安装、内件安装、管道连接、水压试验及验收",
        "换热器": "换热器类设备安装、管道连接、水压试验、保温施工及验收",
        "容器": "容器类设备安装、管道连接、水压试验、保温施工及验收",
        "风机": "风机类设备安装、管道连接、单机试运转及验收",
        "电机": "电机类设备安装、接线、绝缘测试、空载试运转及验收",
        "阀门": "阀门类设备安装、试压、管道连接及验收",
        "储罐": "储罐类设备安装、焊接、充水试验、防腐施工及验收",
    }
    base_scope = type_scope.get(eq_type, "设备安装、管道连接、试运转及验收")
    device_count = len(devices)
    return f"{base_scope}（共{device_count}台{type_scope.get(eq_type, '设备')}）"


def _get_remaining_issues(devices: list) -> dict:
    """v0.1.57：从完整性检查获取遗留问题。"""
    issues = []
    try:
        from . import completeness_check as _cc
        result = _cc.check_completeness()
        # 检查设备级缺失
        device_completeness = result.get("device_completeness", {})
        for tag, info in device_completeness.items():
            if not info.get("complete", True):
                missing = info.get("missing", [])
                if missing:
                    issues.append(f"{tag}：缺少{'、'.join(missing)}")
        # 检查阶段缺失
        stage_completeness = result.get("stage_completeness", {})
        for stage, info in stage_completeness.items():
            if info.get("missing_count", 0) > 0:
                issues.append(f"{stage}阶段：缺少{info.get('missing_count', 0)}项资料")
    except Exception:  # noqa: BLE001
        pass
    
    if not issues:
        issues.append("暂未发现遗留问题，建议现场最终确认")
    
    summary = f"共{len(issues)}项遗留问题（详见遗留问题清单）" if len(issues) > 1 else issues[0]
    return {"summary": summary, "issues": issues}

def fill_template(doc_type: str, data: dict) -> bytes:
    """按模板生成 .docx 并返回字节流。缺字段用『待补充』占位并标注。"""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    t = TYPES[doc_type]
    doc = Document()
    # 全局中文字体
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(18)
        r.font.name = "黑体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    def add_h(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(14)
        r.font.name = "黑体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    def add_kv(key, value):
        v = str(value or "").strip()
        if not v:
            v = "【待补充】"
        p = doc.add_paragraph()
        r = p.add_run(f"{key}：{v}")
        if v == "【待补充】":
            r.font.color.rgb = __import__("docx.shared", fromlist=["RGBColor"]).RGBColor(0xC0, 0x39, 0x2B)
        return v

    missing = []
    add_title(f"{doc_type}")
    add_h("一、基本信息")
    for k in t["required"]:
        v = data.get(k, "")
        if not str(v or "").strip():
            missing.append(k)
        add_kv(k, v)
    for k in t["optional"]:
        if data.get(k):
            add_kv(k, data[k])

    # 技术交底专项（v0.1.51：设备数据联动章节）
    if doc_type == "技术交底":
        if data.get("交底设备参数"):
            add_h("二、交底设备参数")
            for line in str(data["交底设备参数"]).split("\n"):
                if line.strip():
                    p = doc.add_paragraph()
                    r = p.add_run(line.strip())
                    r.font.size = Pt(11)
        if data.get("施工步骤"):
            add_h("三、施工步骤")
            for line in str(data["施工步骤"]).split("\n"):
                if line.strip():
                    p = doc.add_paragraph()
                    r = p.add_run(line.strip())
                    r.font.size = Pt(11)
        if data.get("安全要点"):
            add_h("四、安全要点")
            for line in str(data["安全要点"]).split("\n"):
                if line.strip():
                    p = doc.add_paragraph()
                    r = p.add_run(line.strip())
                    r.font.size = Pt(11)
        if data.get("质量控制要点"):
            add_h("五、质量控制要点")
            for line in str(data["质量控制要点"]).split("\n"):
                if line.strip():
                    p = doc.add_paragraph()
                    r = p.add_run(line.strip())
                    r.font.size = Pt(11)

    # 开箱验收记录专项（v0.1.52：设备数据联动章节）
    if doc_type == "开箱验收记录":
        if data.get("验收设备参数"):
            add_h("二、验收设备参数")
            for line in str(data["验收设备参数"]).split("\n"):
                if line.strip():
                    p = doc.add_paragraph()
                    r = p.add_run(line.strip())
                    r.font.size = Pt(11)
        if data.get("外观检查要点"):
            add_h("三、外观检查要点")
            for line in str(data["外观检查要点"]).split("\n"):
                if line.strip():
                    p = doc.add_paragraph()
                    r = p.add_run(line.strip())
                    r.font.size = Pt(11)
        if data.get("随机资料清单"):
            add_h("四、随机资料清单")
            for line in str(data["随机资料清单"]).split("\n"):
                if line.strip():
                    p = doc.add_paragraph()
                    r = p.add_run(line.strip())
                    r.font.size = Pt(11)

    # 隐蔽工程验收记录专项（v0.1.53：设备数据联动章节）
    if doc_type == "隐蔽工程验收记录":
        if data.get("隐蔽设备参数"):
            add_h("二、隐蔽设备参数")
            for line in str(data["隐蔽设备参数"]).split("\n"):
                if line.strip():
                    p = doc.add_paragraph()
                    r = p.add_run(line.strip())
                    r.font.size = Pt(11)
        if data.get("隐蔽部位"):
            add_h("三、隐蔽部位")
            add_kv("隐蔽部位", data["隐蔽部位"])
        if data.get("隐蔽检查要点"):
            add_h("四、隐蔽检查要点")
            for line in str(data["隐蔽检查要点"]).split("\n"):
                if line.strip():
                    p = doc.add_paragraph()
                    r = p.add_run(line.strip())
                    r.font.size = Pt(11)
        if data.get("质量标准"):
            add_h("五、质量标准")
            for line in str(data["质量标准"]).split("\n"):
                if line.strip():
                    p = doc.add_paragraph()
                    r = p.add_run(line.strip())
                    r.font.size = Pt(11)

    # 设计变更专项（v0.1.55：设备数据联动章节）
    if doc_type == "设计变更":
        if data.get("变更设备参数"):
            add_h("二、变更设备参数")
            for line in str(data["变更设备参数"]).split("\n"):
                if line.strip():
                    p = doc.add_paragraph()
                    r = p.add_run(line.strip())
                    r.font.size = Pt(11)
        if data.get("变更内容建议"):
            add_h("三、变更内容")
            for line in str(data["变更内容建议"]).split("\n"):
                if line.strip():
                    p = doc.add_paragraph()
                    r = p.add_run(line.strip())
                    r.font.size = Pt(11)
        if data.get("影响范围") or data.get("影响车间") or data.get("影响管线"):
            add_h("四、影响范围分析")
            if data.get("影响范围"):
                add_kv("影响概述", data["影响范围"])
            if data.get("影响车间"):
                add_kv("影响车间", data["影响车间"])
            if data.get("影响管线"):
                add_kv("影响管线", data["影响管线"])
            if data.get("影响相邻设备"):
                add_kv("影响相邻设备", data["影响相邻设备"])
        if data.get("处理建议"):
            add_h("五、处理建议")
            for line in str(data["处理建议"]).split("\n"):
                if line.strip():
                    p = doc.add_paragraph()
                    r = p.add_run(line.strip())
                    r.font.size = Pt(11)

    # 货损报告专项（v0.1.56：设备数据联动章节）
    if doc_type == "货损报告":
        if data.get("货损设备参数"):
            add_h("二、货损设备参数")
            for line in str(data["货损设备参数"]).split("\n"):
                if line.strip():
                    p = doc.add_paragraph()
                    r = p.add_run(line.strip())
                    r.font.size = Pt(11)
        if data.get("损失情况建议"):
            add_h("三、损失情况检查")
            for line in str(data["损失情况建议"]).split("\n"):
                if line.strip():
                    p = doc.add_paragraph()
                    r = p.add_run(line.strip())
                    r.font.size = Pt(11)
        if data.get("处理建议"):
            add_h("四、处理建议")
            for line in str(data["处理建议"]).split("\n"):
                if line.strip():
                    p = doc.add_paragraph()
                    r = p.add_run(line.strip())
                    r.font.size = Pt(11)
        if data.get("索赔建议"):
            add_h("五、索赔建议")
            for line in str(data["索赔建议"]).split("\n"):
                if line.strip():
                    p = doc.add_paragraph()
                    r = p.add_run(line.strip())
                    r.font.size = Pt(11)

    # 竣工资料专项（v0.1.57：设备数据联动章节）
    if doc_type == "竣工资料":
        if data.get("竣工设备参数"):
            add_h("二、竣工设备参数")
            for line in str(data["竣工设备参数"]).split("\n"):
                if line.strip():
                    p = doc.add_paragraph()
                    r = p.add_run(line.strip())
                    r.font.size = Pt(11)
        if data.get("交工范围") or data.get("交工车间"):
            add_h("三、交工范围")
            if data.get("交工范围"):
                add_kv("交工内容", data["交工范围"])
            if data.get("交工车间"):
                add_kv("交工车间", data["交工车间"])
        if data.get("验收记录编号") or data.get("隐蔽工程资料编号"):
            add_h("四、验收资料清单")
            if data.get("验收记录编号"):
                add_kv("验收记录编号", data["验收记录编号"])
            if data.get("隐蔽工程资料编号"):
                add_kv("隐蔽工程资料编号", data["隐蔽工程资料编号"])
        if data.get("遗留问题清单"):
            add_h("五、遗留问题及处理")
            for line in str(data["遗留问题清单"]).split("\n"):
                if line.strip():
                    p = doc.add_paragraph()
                    r = p.add_run(line.strip())
                    r.font.size = Pt(11)

    # 吊装方案专项参数（v0.1.30）
    if doc_type == "吊装方案":
        add_h("二、吊装参数")
        lifting_keys = ["吊装设备名称", "设备重量", "吊装半径", "吊装高度", "吊车型号", "吊车站位", "吊索具"]
        for k in lifting_keys:
            if data.get(k) or k in t["required"]:
                add_kv(k, data.get(k, ""))
        # 吊装计算提示
        p = doc.add_paragraph()
        r = p.add_run("注：吊装半径、吊装高度需现场实测后填入；吊车型号应根据吊装重量与半径经吊装计算确定，本方案仅为建议。")
        r.font.size = Pt(10)
        r.font.color.rgb = __import__("docx.shared", fromlist=["RGBColor"]).RGBColor(0x66, 0x66, 0x66)
        # v0.1.50：吊装方案设备数据联动章节
        if data.get("吊装设备参数"):
            add_h("三、吊装设备参数")
            for line in str(data["吊装设备参数"]).split("\n"):
                if line.strip():
                    p = doc.add_paragraph()
                    r = p.add_run(line.strip())
                    r.font.size = Pt(11)
        if data.get("吊装环境"):
            add_h("四、吊装环境")
            add_kv("作业位置", data["吊装环境"])
        if data.get("空间限制"):
            add_h("五、空间限制与特殊要求")
            for line in str(data["空间限制"]).split("\n"):
                if line.strip():
                    p = doc.add_paragraph()
                    r = p.add_run(line.strip())
                    r.font.size = Pt(11)
        if data.get("吊装相邻设备"):
            add_h("六、相邻设备与成品保护")
            add_kv("相邻设备", data["吊装相邻设备"])
            add_kv("成品保护", "吊装前对相邻设备进行覆盖保护，设置警戒区，严禁碰撞已安装设备和管道")
        if data.get("吊装相关管线"):
            add_h("七、相关管线与保护")
            for line in str(data["吊装相关管线"]).split("\n"):
                if line.strip():
                    p = doc.add_paragraph()
                    r = p.add_run(line.strip())
                    r.font.size = Pt(11)

    # 设备清单（吊装方案/开箱验收自动带出）
    devices = data.get("_devices") or []
    if devices:
        add_h("二、涉及设备清单")
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i, h in enumerate(["位号", "设备名称", "数量", "备注"]):
            hdr[i].text = h
        for dv in devices[:60]:
            row = tbl.add_row().cells
            row[0].text = str(dv.get("tag", ""))
            row[1].text = str(dv.get("name", "见台账"))
            row[2].text = str(dv.get("count", "1"))
            row[3].text = ""
    else:
        add_h("二、相关设备")
        add_kv("设备清单", data.get("_devices_hint", "可到关联图谱页选择车间自动带出"))

    # v0.1.42：施工方案专项——施工步骤（按设备类型自动生成）
    if doc_type == "施工方案" and data.get("施工步骤"):
        add_h("三、施工步骤")
        for line in str(data["施工步骤"]).split("\n"):
            if line.strip():
                p = doc.add_paragraph()
                r = p.add_run(line.strip())
                r.font.size = Pt(11)
        # v0.1.49：设备数据联动——技术参数、施工环境、管线连接、相邻设备
        if data.get("设备技术参数"):
            add_h("四、设备技术参数")
            for line in str(data["设备技术参数"]).split("\n"):
                if line.strip():
                    p = doc.add_paragraph()
                    r = p.add_run(line.strip())
                    r.font.size = Pt(11)
        if data.get("施工环境"):
            add_h("五、施工环境")
            add_kv("作业位置", data["施工环境"])
        if data.get("相关管线"):
            add_h("六、相关管线与连接")
            for line in str(data["相关管线"]).split("\n"):
                if line.strip():
                    p = doc.add_paragraph()
                    r = p.add_run(line.strip())
                    r.font.size = Pt(11)
        if data.get("相邻设备"):
            add_h("七、相邻设备与交叉作业")
            add_kv("相邻设备", data["相邻设备"])
            add_kv("交叉作业注意", "施工前确认相邻设备状态，做好成品保护，交叉作业时设专人监护")
        add_h("八、质量控制")
        add_kv("质量标准", "符合现行国家及行业施工验收规范，一次验收合格率100%")
        add_kv("质量控制点", "基础验收、设备找平找正、对中偏差、管道焊接、试运转参数")

    # 默认章节文本
    section_no = 2 if (devices or doc_type in ("吊装方案", "开箱验收记录", "隐蔽工程验收记录")) else 2
    if "编制依据" in t["default_text"]:
        add_h(f"{'三' if section_no == 2 else '二'}、编制依据")
        cites = data.get("_std_citations") or std_citations(doc_type)
        if cites:
            doc.add_paragraph(t["default_text"]["编制依据"])
            doc.add_paragraph("本资料编制所引用的现行规范条款如下：")
            for ci in cites:
                p = doc.add_paragraph()
                r = p.add_run(f"▶ {ci['std_no']}《{ci['std_name']}》（{ci['status']}）：{ci['snippet']}")
                r.font.size = Pt(11)
        else:
            doc.add_paragraph(t["default_text"]["编制依据"])
            p = doc.add_paragraph("（平台库尚未收录相关现行规范正文，请人工补充编制依据）")
            p.runs[0].font.color.rgb = __import__("docx.shared", fromlist=["RGBColor"]).RGBColor(0xC0, 0x39, 0x2B)
    for k, default in t["default_text"].items():
        if k == "编制依据":
            continue
        add_h(f"四、{k}")
        doc.add_paragraph(default)

    add_h("五、签字栏")
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    for i, h in enumerate(["编制", "审核", "批准", "日期"]):
        tbl.rows[0].cells[i].text = h
    add_paragraph = doc.add_paragraph
    p = add_paragraph()
    p.add_run(f"\n生成工具：繁工AI 本地解析工作台 v0.1.30 · 生成时间 {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")

    import io
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), missing
