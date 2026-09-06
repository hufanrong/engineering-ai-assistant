# 繁工AI 本地解析工作台 - 资料自动关联到设备/车间（v0.1.39）
# 目的：已生成的工程资料（施工方案、开箱验收、隐蔽记录等）自动关联到
#       对应的设备和车间，让完整性检查更精准，AI 能知道哪些设备已有资料。
#
# 关联来源（按优先级）：
#   1. 生成时传入的 metadata（docgen generate 时的设备/车间）
#   2. 文件名中的位号/车间（如 繁工AI_开箱验收记录_P-101_1号车间.docx）
#   3. 资料内容中的位号/车间（正则扫描 docx 文本）
#   4. 已上传资料文件的文件名匹配

import os
import re
import json
import datetime

from . import config


_RELATIONS_FILE = os.path.join(config.DATA_DIR, "doc_relations.json")

# 位号正则：P-101, V-201, E-101A, T-301 等
_TAG_RE = re.compile(r"(?<![A-Z0-9])([A-Z]{1,4}-\d{2,4}[A-Z]?)(?![A-Z0-9])")
# 车间正则：1号车间、2#车间、一号车间、车间1、A车间 等
_WORKSHOP_RE = re.compile(
    r"(\d+号车间|\d+#车间|[一二三四五六七八九十]+号车间|车间\d+|[A-Z]车间|[A-Z]区)"
)


def _load() -> dict:
    if os.path.exists(_RELATIONS_FILE):
        try:
            with open(_RELATIONS_FILE, encoding="utf-8") as f:
                return json.load(f)
        except Exception:  # noqa: BLE001
            return {}
    return {}


def _save(data: dict):
    os.makedirs(os.path.dirname(_RELATIONS_FILE), exist_ok=True)
    with open(_RELATIONS_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def extract_from_text(text: str) -> dict:
    """从文本中提取位号和车间。"""
    tags = set()
    workshops = set()
    if not text:
        return {"tags": [], "workshops": []}
    for m in _TAG_RE.finditer(text):
        tag = m.group(1)
        # 过滤掉明显不是位号的（如日期、版本号）
        if not re.match(r"^(V|E|P|T|R|C|K|F|Y|Z|Q|S|A|B|D|G|H|J|L|M|N|O|U|W|X)-", tag, re.IGNORECASE):
            # 宽松匹配：只要是 字母-数字 格式就保留
            if re.match(r"^[A-Z]+-\d+", tag):
                tags.add(tag)
        else:
            tags.add(tag)
    for m in _WORKSHOP_RE.finditer(text):
        workshops.add(m.group(1))
    return {"tags": sorted(tags), "workshops": sorted(workshops)}


def extract_from_filename(filename: str) -> dict:
    """从文件名中提取位号和车间。"""
    return extract_from_text(filename)


def extract_from_docx(path: str) -> dict:
    """从 docx 文件中提取位号和车间（扫描段落文本）。"""
    try:
        from docx import Document
        doc = Document(path)
        text_parts = []
        for p in doc.paragraphs:
            if p.text:
                text_parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text_parts.append(cell.text)
        full_text = "\n".join(text_parts)
        return extract_from_text(full_text)
    except Exception:  # noqa: BLE001
        return {"tags": [], "workshops": []}


def register_doc(doc_id: str, doc_type: str, file_path: str = "",
                 devices: list = None, workshops: list = None,
                 source: str = "generated") -> dict:
    """登记一份资料的设备/车间关联。

    Args:
        doc_id: 资料唯一标识（文件名或 sha）
        doc_type: 资料类型（施工方案、开箱验收记录等）
        file_path: 文件路径（用于从内容提取）
        devices: 已知关联设备列表
        workshops: 已知关联车间列表
        source: 来源（generated/uploaded）

    Returns:
        关联记录 dict
    """
    data = _load()
    record = {
        "doc_id": doc_id,
        "doc_type": doc_type,
        "file_path": file_path,
        "source": source,
        "devices": list(set(devices or [])),
        "workshops": list(set(workshops or [])),
        "registered_at": datetime.datetime.now().isoformat(),
    }

    # 从文件名补充
    if file_path:
        fname = os.path.basename(file_path)
        fn_extract = extract_from_filename(fname)
        for t in fn_extract["tags"]:
            if t not in record["devices"]:
                record["devices"].append(t)
        for w in fn_extract["workshops"]:
            if w not in record["workshops"]:
                record["workshops"].append(w)

        # 从 docx 内容补充（如果文件存在且是 docx）
        if file_path.endswith(".docx") and os.path.exists(file_path):
            content_extract = extract_from_docx(file_path)
            for t in content_extract["tags"]:
                if t not in record["devices"]:
                    record["devices"].append(t)
            for w in content_extract["workshops"]:
                if w not in record["workshops"]:
                    record["workshops"].append(w)

    data[doc_id] = record
    _save(data)
    return record


def scan_generated_docs() -> int:
    """扫描 data/generated_docs/ 下所有已生成资料，自动登记关联。"""
    from . import archive
    gen_dir = archive._gen_dir()
    if not os.path.isdir(gen_dir):
        return 0
    count = 0
    for doc_type in os.listdir(gen_dir):
        tdir = os.path.join(gen_dir, doc_type)
        if not os.path.isdir(tdir):
            continue
        for fname in os.listdir(tdir):
            if not fname.endswith(".docx"):
                continue
            fpath = os.path.join(tdir, fname)
            doc_id = f"{doc_type}/{fname}"
            register_doc(doc_id, doc_type, file_path=fpath, source="generated")
            count += 1
    return count


def get_by_device(tag: str) -> list:
    """获取某台设备关联的所有资料。"""
    data = _load()
    results = []
    for doc_id, rec in data.items():
        if tag in rec.get("devices", []):
            results.append(rec)
    return results


def get_by_workshop(workshop: str) -> list:
    """获取某个车间关联的所有资料。"""
    data = _load()
    results = []
    for doc_id, rec in data.items():
        if workshop in rec.get("workshops", []):
            results.append(rec)
    return results


def get_all() -> dict:
    """获取所有资料关联。"""
    return _load()


def stats() -> dict:
    """关联统计。"""
    data = _load()
    device_docs = {}
    workshop_docs = {}
    type_count = {}
    for doc_id, rec in data.items():
        dtype = rec.get("doc_type", "未知")
        type_count[dtype] = type_count.get(dtype, 0) + 1
        for dev in rec.get("devices", []):
            device_docs[dev] = device_docs.get(dev, 0) + 1
        for ws in rec.get("workshops", []):
            workshop_docs[ws] = workshop_docs.get(ws, 0) + 1
    return {
        "total_docs": len(data),
        "with_devices": sum(1 for r in data.values() if r.get("devices")),
        "with_workshops": sum(1 for r in data.values() if r.get("workshops")),
        "type_count": type_count,
        "device_docs": device_docs,
        "workshop_docs": workshop_docs,
    }


def device_has_doc_type(tag: str, doc_type: str) -> bool:
    """检查某台设备是否已有指定类型的资料。"""
    docs = get_by_device(tag)
    return any(d.get("doc_type") == doc_type for d in docs)


def workshop_has_doc_type(workshop: str, doc_type: str) -> bool:
    """检查某个车间是否已有指定类型的资料。"""
    docs = get_by_workshop(workshop)
    return any(d.get("doc_type") == doc_type for d in docs)
