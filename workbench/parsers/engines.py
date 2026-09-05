# 繁工AI 本地解析工作台 - 深度解析引擎
# 设计：按文件类型分派解析器；每个解析器输出统一结构 ParseResult；
# 可选依赖（OCR/CAD）缺失时优雅降级，不影响主链路。
# 优先级（v3.6 口径）：Project → 表格台账 → 图片OCR → CAD图纸

import os
import re
import json
import hashlib
import datetime
from dataclasses import dataclass, field, asdict
from typing import Optional

from app import config


# ============ 统一结果结构 ============
@dataclass
class ParseResult:
    file_path: str                 # 原始路径
    file_name: str                 # 文件名
    file_size: int                 # 字节
    sha256: str                    # 文件哈希（去重/云端合并用）
    ext: str                       # 扩展名
    parser: str                    # 解析器名
    status: str = "parsed"         # parsed / partial / skipped / failed
    error: Optional[str] = None    # 失败原因
    text: str = ""                 # 提取的全文（用于向量化）
    structure: dict = field(default_factory=dict)   # 结构化结果（表头/行/任务树/标题栏等）
    entities: list = field(default_factory=list)    # 浅层实体（设备位号等）
    chunks: list = field(default_factory=list)      # 分块（由向量化步骤填充）
    created_at: str = field(default_factory=lambda: datetime.datetime.now().isoformat())


def sha256_of_file(path: str) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for block in iter(lambda: f.read(65536), b""):
            h.update(block)
    return h.hexdigest()


# ============ 实体浅提取（位号/编号） ============
_TAG_RE = re.compile(config.EQUIPMENT_TAG_RE)


_PARAM_RE = re.compile(
    r"(?:流量|扬程|功率|电压|电流|转速|重量|介质|温度|压力|频率|防爆等级|防护等级|绝缘等级|出厂编号|出厂日期|重量|容积|材质)"
    r"\s*[:：]?\s*[0-9A-Za-z./×xX·\-~～%℃m³kgKVWkWh]*")
_MFR_RE = re.compile(r"[\u4e00-\u9fa5A-Za-z0-9（）()·]{2,20}?(?:有限公司|股份公司|制造厂|机械厂|实业公司|集团)")


def extract_entities(text: str, limit: int = 200) -> list:
    """从文本中提取实体：设备位号 / 铭牌参数 / 厂家 / 车间。
    浅层提取，后续由云端图谱归并去重。"""
    seen, out = set(), []
    for m in _TAG_RE.finditer(text or ""):
        tag = m.group(1)
        if tag not in seen:
            seen.add(tag)
            out.append({"type": "equipment", "tag": tag})
        if len(out) >= limit:
            break
    if len(out) < limit:
        for m in _PARAM_RE.finditer(text or ""):
            kv = m.group(0).strip()
            key = kv.split("：")[0].split(":")[0].strip()
            if kv not in seen:
                seen.add(kv)
                out.append({"type": "parameter", "name": key, "value": kv})
            if len(out) >= limit:
                break
    if len(out) < limit:
        for m in _MFR_RE.finditer(text or ""):
            v = m.group(0)
            if v not in seen:
                seen.add(v)
                out.append({"type": "manufacturer", "name": v})
            if len(out) >= limit:
                break
    if len(out) < limit:
        for m in re.finditer(r"\d+\s*号\s*车间", text or ""):
            v = m.group(0)
            if v not in seen:
                seen.add(v)
                out.append({"type": "workshop", "name": v})
            if len(out) >= limit:
                break
    return out


# ============ 解析器注册表（优先级：Project → 表格台账 → 图片OCR → CAD → PDF/Word/Text） ============
def _effective(switch: bool, dep: str) -> bool:
    """有效开关：自动探测开启且依赖已装 → 直接启用（全套部署免配置）；
    否则按 config 里的开关。"""
    if config.AUTO_DETECT_OPTIONAL and config.OPTIONAL_READY.get(dep, False):
        return True
    return switch


def _pick_parser(ext: str):
    if ext in config.EXT_PROJECT and config.PARSE_PROJECT:
        return parse_project
    if ext in config.EXT_EXCEL and config.PARSE_EXCEL:
        return parse_excel
    if ext in config.EXT_IMAGE and _effective(config.PARSE_IMAGE, "ocr"):
        return parse_image
    if ext in config.EXT_CAD and _effective(config.PARSE_CAD, "cad"):
        return parse_cad
    if ext in config.EXT_PDF and config.PARSE_PDF:
        return parse_pdf
    if ext in config.EXT_WORD and config.PARSE_WORD:
        return parse_word
    if ext in config.EXT_TEXT and config.PARSE_TEXT:
        return parse_text
    return None


def parse_file(path: str) -> ParseResult:
    """按扩展名分派解析。无法解析的类型返回 skipped。"""
    ext = os.path.splitext(path)[1].lower()
    name = os.path.basename(path)
    size = os.path.getsize(path)
    res = ParseResult(
        file_path=path, file_name=name, file_size=size,
        sha256=sha256_of_file(path), ext=ext, parser="none",
    )
    parser = _pick_parser(ext)
    if parser is None:
        res.status = "skipped"
        res.parser = "unsupported"
        res.error = "该类型当前未启用解析（可在 config.py 开启或安装可选依赖）"
        return res
    try:
        parser(res)
        if not res.text and not res.structure:
            res.status = "partial"
        res.entities = extract_entities(res.text)
    except ImportError as e:
        res.status = "skipped"
        res.error = f"缺少依赖：{e}。请按 README 安装可选依赖后重启"
    except Exception as e:  # noqa: BLE001
        res.status = "failed"
        res.error = f"{type(e).__name__}: {e}"
    return res


# ============ PDF（文本 + 表格；pdfplumber 缺失时降级为纯文本） ============
def parse_pdf(res: ParseResult):
    from pypdf import PdfReader
    res.parser = "pdf"
    reader = PdfReader(res.file_path)
    parts = []
    for i, page in enumerate(reader.pages):
        try:
            t = page.extract_text() or ""
        except Exception:  # noqa: BLE001
            t = ""
        if t.strip():
            parts.append(f"【第{i+1}页】\n{t}")
    res.text = "\n\n".join(parts)
    tables = _extract_pdf_tables(res.file_path)
    if tables:
        for ti, tb in enumerate(tables):
            parts.append(f"【PDF表格{ti+1}】\n" + "\n".join(" | ".join(row) for row in tb))
        res.text = "\n\n".join(parts)
    res.structure = {
        "page_count": len(reader.pages),
        "text_chars": len(res.text),
        "tables": tables[:50],
    }
    if not res.text:
        res.status = "partial"
        res.error = "PDF 无可提取文本层（可能是扫描件，建议启用 OCR）"


def _extract_pdf_tables(path: str) -> list:
    """用 pdfplumber 提取跨页表格（可选依赖；缺失/失败时返回空，不影响主流程）。
    先按画线识别；无边框表格（常见于导出的 PDF）再用文本策略兜底。"""
    try:
        import pdfplumber
    except ImportError:
        return []

    def _run(table_settings=None):
        out = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                for tb in page.extract_tables(table_settings=table_settings):
                    rows = [[(c or "").replace("\n", " ").strip() for c in row] for row in tb]
                    if any(any(c for c in row) for row in rows):
                        out.append(rows)
        return out

    try:
        out = _run()
        if out:
            return out
        # 无画线表格 → 按文本对齐策略识别
        return _run({"vertical_strategy": "text", "horizontal_strategy": "text"})
    except Exception:  # noqa: BLE001
        return []


# ============ Word ============
def parse_word(res: ParseResult):
    import docx
    res.parser = "word"
    doc = docx.Document(res.file_path)
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    tables = []
    for ti, tb in enumerate(doc.tables):
        rows = []
        for row in tb.rows:
            rows.append([c.text.strip() for c in row.cells])
        tables.append({"sheet": f"表格{ti+1}", "rows": rows})
        for row in rows:
            parts.append(" | ".join(row))
    res.text = "\n".join(parts)
    res.structure = {"paragraphs": len(doc.paragraphs), "tables": tables}


# ============ Excel（台账结构化，v3.6 核心） ============
def _normalize_header(cells):
    """表头归一：合并单元格/多行表头 → 取首个非空作为列名；同名列自动加序号去重。"""
    out = []
    seen = {}
    for c in cells:
        v = (c or "").strip()
        name = v if v else f"列{len(out)+1}"
        if name in seen:
            seen[name] += 1
            name = f"{name}_{seen[name]}"
        else:
            seen[name] = 0
        out.append(name)
    return out


def parse_excel(res: ParseResult):
    import openpyxl
    res.parser = "excel"
    wb = openpyxl.load_workbook(res.file_path, read_only=True, data_only=True)
    sheets = []
    text_parts = []
    for ws in wb.worksheets:
        rows_iter = ws.iter_rows(values_only=True)
        all_rows = [list(r) for r in rows_iter if any(v is not None and str(v).strip() for v in r)]
        if not all_rows:
            continue
        header = _normalize_header([str(v) if v is not None else "" for v in all_rows[0]])
        data_rows = []
        for r in all_rows[1:]:
            row = {}
            for i, v in enumerate(r):
                if i < len(header) and v is not None:
                    row[header[i]] = str(v).strip()
            if row:
                data_rows.append(row)
        sheets.append({
            "sheet": ws.title,
            "header": header,
            "row_count": len(data_rows),
            "rows": data_rows[:500],          # 前端展示上限
        })
        text_parts.append(f"【工作表：{ws.title}】")
        text_parts.append(" | ".join(header))
        for row in data_rows[:1000]:
            text_parts.append(" | ".join(f"{k}={v}" for k, v in row.items()))
    wb.close()
    res.text = "\n".join(text_parts)
    res.structure = {"sheets": sheets, "total_rows": sum(s["row_count"] for s in sheets)}


# ============ 文本 ============
def parse_text(res: ParseResult):
    res.parser = "text"
    with open(res.file_path, "r", encoding="utf-8", errors="ignore") as f:
        res.text = f.read()
    res.structure = {"chars": len(res.text)}


# ============ 图片 OCR（可选，需 PaddleOCR） ============
_ocr_engine = None


def _get_ocr():
    global _ocr_engine
    if _ocr_engine is None:
        from paddleocr import PaddleOCR
        _ocr_engine = PaddleOCR(use_angle_cls=True, lang="ch", show_log=False)
    return _ocr_engine


def parse_image(res: ParseResult):
    """图片 OCR：提取文本块（带置信度与位置）→ 铭牌结构化（位号/型号/参数/厂家）→ 全文。
    支持 PaddleOCR 2.x（ocr.ocr）与 3.x（ocr.predict）。"""
    res.parser = "ocr"
    ocr = _get_ocr()
    lines = []
    blocks = []
    try:
        result = ocr.ocr(res.file_path, cls=True)          # PaddleOCR 2.x
        for page in result or []:
            for item in page or []:
                box, (text, conf) = item
                lines.append(text)
                blocks.append({"text": text, "conf": round(float(conf), 3), "box": box})
    except TypeError:
        result = ocr.predict(res.file_path)                # PaddleOCR 3.x
        for page in result or []:
            for item in getattr(page, "rec_texts", []) or []:
                lines.append(item)
                blocks.append({"text": item, "conf": None, "box": None})

    full = "\n".join(lines)
    # 铭牌结构化：位号 / 参数 / 厂家
    plate = {
        "tags": [m.group(1) for m in _TAG_RE.finditer(full)][:20],
        "params": [kv for kv in (_m.group(0).strip() for _m in _PARAM_RE.finditer(full))][:30],
        "manufacturers": [m.group(0) for m in _MFR_RE.finditer(full)][:5],
        "workshops": [m.group(0) for m in re.finditer(r"\d+\s*号\s*车间", full)][:5],
    }
    res.text = full
    res.structure = {
        "blocks": blocks, "line_count": len(lines),
        "plate": plate,
        "is_plate": bool(plate["tags"] or plate["params"] or plate["manufacturers"]),
    }


# ============ CAD（可选，需 ezdxf；DWG 需 ODA 转 DXF）
# 深度解析 v2（v0.1.4）：图框检测 + 标题栏键值提取 + 尺寸标注 + 设备块属性 + 图层统计 + 空间结构
# ============
def parse_cad(res: ParseResult):
    import ezdxf
    res.parser = "cad"
    path = res.file_path
    if path.lower().endswith(".dwg"):
        converted = _dwg_to_dxf(path)
        if not converted:
            res.status = "failed"
            res.error = "DWG 需先安装 ODA File Converter 并配置转换（见 README），或先用 AutoCAD 另存为 DXF"
            return
        path = converted
    doc = ezdxf.readfile(path)
    msp = doc.modelspace()

    texts, labels, tags, dims, layers, lines_p = [], [], [], [], {}, []
    for e in msp:
        t = e.dxftype()
        if t == "TEXT":
            texts.append(e.dxf.text)
            labels.append({"text": e.dxf.text, "x": round(e.dxf.insert.x, 2), "y": round(e.dxf.insert.y, 2), "layer": e.dxf.layer})
        elif t == "MTEXT":
            texts.append(e.plain_text())
            labels.append({"text": e.plain_text(), "x": round(e.dxf.insert.x, 2), "y": round(e.dxf.insert.y, 2), "layer": e.dxf.layer})
        elif t == "INSERT":
            attrs = []
            try:
                for a in e.attribs:
                    attrs.append({"tag": a.dxf.tag, "value": a.dxf.text})
            except Exception:  # noqa: BLE001
                pass
            tags.append({
                "block": e.dxf.name,
                "x": round(e.dxf.insert.x, 2),
                "y": round(e.dxf.insert.y, 2),
                "layer": e.dxf.layer,
                "attrs": attrs[:50],
                "scale": round(e.dxf.xscale or 1, 3),
            })
        elif t == "DIMENSION":
            try:
                m = e.get_measurement()
            except Exception:  # noqa: BLE001
                m = None
            dims.append({
                "text": (e.dxf.text or "").strip(),
                "measurement": round(m, 3) if m is not None else None,
                "x": round(e.dxf.defpoint.x, 2), "y": round(e.dxf.defpoint.y, 2),
                "layer": e.dxf.layer,
            })
        elif t == "LINE":
            lines_p.append((round(e.dxf.start.x, 2), round(e.dxf.start.y, 2), round(e.dxf.end.x, 2), round(e.dxf.end.y, 2)))
        layers[e.dxf.layer] = layers.get(e.dxf.layer, 0) + 1

    # —— 图框检测：取几何范围（若有闭合矩形 LWPOLYLINE 则用其范围）——
    frame = _detect_frame(msp, lines_p)
    # —— 标题栏：图框/图面右下角区域的键值字段提取 ——
    title_fields = _extract_title_block(labels, frame)

    res.text = "\n".join(texts)
    res.structure = {
        "spatial": {
            "frame": frame,                       # 图框边界 [xmin, ymin, xmax, ymax]
            "title_block": title_fields,          # 图号/图名/比例/设计/日期等
            "blocks": tags[:800],                 # 设备/图块 + 坐标 + 属性（空间库打底）
            "dimensions": dims[:800],             # 尺寸标注 + 测量值
        },
        "text_labels": labels[:500],
        "version": doc.dxfversion,
        "layers": [{"layer": k, "count": v} for k, v in sorted(layers.items(), key=lambda x: -x[1])][:100],
        "entity_counts": {
            "TEXT": len([1 for e in msp.query("TEXT")]),
            "MTEXT": len([1 for e in msp.query("MTEXT")]),
            "INSERT": len([1 for e in msp.query("INSERT")]),
            "LINE": len([1 for e in msp.query("LINE")]),
            "DIMENSION": len([1 for e in msp.query("DIMENSION")]),
        },
    }
    res.entities = extract_entities("\n".join(texts + [a.get("value", "") for t in tags for a in t.get("attrs", [])]))


def _detect_frame(msp, lines_p):
    """图框检测：优先找面积最大的闭合 LWPOLYLINE 矩形；否则退回全图几何范围。"""
    best = None
    for e in msp.query("LWPOLYLINE"):
        try:
            pts = list(e.get_points("xy"))
            if len(pts) >= 4 and pts[0] == pts[-1]:
                xs = [p[0] for p in pts]; ys = [p[1] for p in pts]
                w, h = max(xs) - min(xs), max(ys) - min(ys)
                if w > 0 and h > 0:
                    area = w * h
                    if best is None or area > best[0]:
                        best = (area, [round(min(xs), 2), round(min(ys), 2), round(max(xs), 2), round(max(ys), 2)])
        except Exception:  # noqa: BLE001
            continue
    if best:
        return best[1]
    if lines_p:
        xs = [c for p in lines_p for c in (p[0], p[2])]
        ys = [c for p in lines_p for c in (p[1], p[3])]
        return [round(min(xs), 2), round(min(ys), 2), round(max(xs), 2), round(max(ys), 2)]
    return None


_TITLE_KEYS = [
    ("图号", ["图号", "图纸编号", "编号", "drawing no", "drw no"]),
    ("图名", ["图名", "图纸名称", "名称", "title"]),
    ("工程名", ["工程名称", "项目名称", "project"]),
    ("车间", ["车间", "workshop", "unit"]),
    ("专业", ["专业", "discipline"]),
    ("比例", ["比例", "scale"]),
    ("设计", ["设计", "design"]),
    ("制图", ["制图", "drawn", "drawing by"]),
    ("审核", ["审核", "checked", "check"]),
    ("批准", ["批准", "approved"]),
    ("日期", ["日期", "date"]),
    ("版本", ["版本", "rev", "version"]),
    ("图幅", ["图幅", "a0", "a1", "a2", "a3", "a4"]),
]


def _extract_title_block(labels, frame):
    """标题栏键值提取：取图框右下 40% 区域（含右下 20% 高的横条）内的文本，按常见键词解析。
    无图框时用全图文本边界兜底（部分简图/示意图没有图框线）。"""
    if not labels:
        return {}
    if frame is None:
        # 兜底：以文本范围为边界，取右下 50%×50% 区域
        xs = [lb["x"] for lb in labels]
        ys = [lb["y"] for lb in labels]
        frame = [min(xs), min(ys), max(xs), max(ys)]
        xmin, ymin, xmax, ymax = frame
        w, h = xmax - xmin, ymax - ymin
        rx0 = round(xmin + w * 0.5, 2)
        ymax_limit = round(ymin + h * 0.5, 2)
    else:
        xmin, ymin, xmax, ymax = frame
        w, h = xmax - xmin, ymax - ymin
        if w <= 0 or h <= 0:
            return {}
        rx0 = round(xmin + w * 0.55, 2)
        ymax_limit = round(ymin + h * 0.45, 2)
    zone = [lb for lb in labels
            if lb["x"] >= rx0 and lb["y"] >= ymin
            and lb["x"] <= xmax and lb["y"] <= ymax_limit]
    if not zone:
        return {}
    # 同一行按 y 聚拢，按 x 排序拼成行文本
    rows = {}
    for lb in zone:
        key = round(lb["y"], 1)
        rows.setdefault(key, []).append(lb)
    row_texts = []
    for key in sorted(rows.keys(), reverse=True):   # 从高到低（标题栏顶部行优先）
        row_texts.append(" ".join(lb["text"] for lb in sorted(rows[key], key=lambda z: z["x"])))

    out = {}
    joined = "\n".join(row_texts)
    for field, keys in _TITLE_KEYS:
        if field in out:
            continue
        for k in keys:
            # 独立键匹配："键 值" 或 "键:值"，键前必须是行首或分隔符（避免命中“1号车间”里的“车间”）
            m = re.search(r"(?:^|[\s:：])({})[\s:：]*(.+)$".format(re.escape(k)), joined, re.M)
            if not m:
                continue
            rest = m.group(2).strip()
            if rest:
                out[field] = rest
                break
    # 若整区无键词，取最大字体/首行作为图名候选
    if not out and zone:
        zone_sorted = sorted(zone, key=lambda z: -z["y"])
        out["图名(候选)"] = zone_sorted[0]["text"]
    return out


def _dwg_to_dxf(dwg_path: str) -> Optional[str]:
    """调用 ODA File Converter 命令行将 DWG 转 DXF；失败返回 None。"""
    import subprocess
    oda = os.environ.get("ODA_CONVERTER", r"C:\Program Files\ODA\ODAFileConverter\ODAFileConverter.exe")
    if not os.path.exists(oda):
        return None
    src_dir = os.path.dirname(dwg_path)
    out_dir = os.path.join(src_dir, "_dxf_converted")
    os.makedirs(out_dir, exist_ok=True)
    # ODAFileConverter 输入目录 输出目录 输出版本 输出类型 递归 审计
    cmd = [oda, src_dir, out_dir, "ACAD2018", "DXF", "0", "1"]
    try:
        subprocess.run(cmd, timeout=180, check=True, capture_output=True)
        base = os.path.splitext(os.path.basename(dwg_path))[0] + ".dxf"
        cand = os.path.join(out_dir, base)
        return cand if os.path.exists(cand) else None
    except Exception:  # noqa: BLE001
        return None


# ============ Project 计划（.xml，纯 Python；.mpp 需先另存为 XML） ============
def _project_link_uid(task_elem, ns) -> str:
    """取 PredecessorLink 下的 PredecessorUID 子元素（前置任务编号）。"""
    link = task_elem.find("p:PredecessorLink", ns)
    if link is None:
        return ""
    uids = [e.text or "" for e in link.findall("p:PredecessorUID", ns) if e.text]
    return ",".join(uids)


def parse_project(res: ParseResult):
    import xml.etree.ElementTree as ET
    res.parser = "project"
    tree = ET.parse(res.file_path)
    root = tree.getroot()
    ns = {"p": "http://schemas.microsoft.com/project"}

    # 任务：名称 / WBS / 开始完成 / 工期 / 前置任务 / 里程碑
    tasks = []
    for t in root.findall(".//p:Task", ns):
        def g(tag):
            e = t.find(f"p:{tag}", ns)
            return (e.text or "").strip() if e is not None and e.text else ""
        task = {
            "id": g("UID"),
            "name": g("Name"),
            "wbs": g("WBS"),
            "start": g("Start"),
            "finish": g("Finish"),
            "duration": g("Duration"),
            "predecessors": g("PredecessorLink") or g("Predecessors")
                or _project_link_uid(t, ns),
            "milestone": g("Milestone"),
        }
        if task["name"]:
            tasks.append(task)

    text_parts = []
    for t in tasks:
        line = " | ".join(f"{k}={v}" for k, v in t.items() if v)
        text_parts.append(line)
    res.text = "\n".join(text_parts)
    res.structure = {"tasks": tasks, "task_count": len(tasks)}
    if not tasks:
        res.status = "partial"
        res.error = "未解析到任务，确认文件是 MS Project 导出的 XML（文件头含 <Project>）"
